import os
import concurrent.futures
from typing import Dict, List, Optional, Tuple

from docx import Document
from .schemas import OutlineResponse
from .transcript_utils import (
    get_playlist_info,
    get_video_info,
    get_video_transcript,
    captions_to_plaintext,
    captions_to_segments,
)
from .llm_utils import outline_for_text, merge_outlines


"""
High-level orchestration functions for generating DOCX notes for playlists and single videos.
Transcript fetching/parsing and LLM calls are delegated to dedicated modules.
"""


def generate_playlist_docx(
    playlist_url: str,
    output_dir: str = "playlists_docx",
    *,
    use_llm: bool = True,
    llm_model: str = "gpt-4o-mini",
    include_raw_transcript: bool = False,
    openai_api_key: Optional[str] = None,
    max_workers: int = 250,
) -> str:
    """
    Create a DOCX file for the provided playlist URL.
    Writes the playlist title and per-video titles with transcript text.
    Returns the saved DOCX file path.
    """
    os.makedirs(output_dir, exist_ok=True)

    playlist_info = get_playlist_info(playlist_url)
    playlist_title = playlist_info.get('title', 'Playlist')

    document = Document()
    document.add_heading(playlist_title, level=1)

    entries = [e for e in (playlist_info.get('entries', []) or []) if e]

    def process_entry(entry: Dict) -> Tuple[str, Optional[OutlineResponse], Optional[str]]:
        video_title = entry.get('title', 'Untitled')
        video_id = entry.get('id')
        if not video_id:
            return (video_title, None, "Missing video id")
        video_url = f"https://www.youtube.com/watch?v={video_id}"
        transcript_text = get_video_transcript(video_url)
        plain = captions_to_plaintext(transcript_text)
        if not use_llm:
            return (video_title, OutlineResponse(sections=[]), None)
        try:
            info = get_video_info(video_url)
            duration = int(info.get('duration') or 0)
            if duration > 1800:
                segments = captions_to_segments(transcript_text, segment_seconds=1800)
                outlines: List[OutlineResponse] = []
                for seg_index, start_sec, seg_text in segments:
                    outline_seg = outline_for_text(
                        f"{video_title} — Part {seg_index+1}",
                        seg_text,
                        model=llm_model,
                        api_key=openai_api_key,
                        playlist_mode=True,
                    )
                    outlines.append(outline_seg)
                outline = merge_outlines(video_title, outlines, model=llm_model, api_key=openai_api_key, playlist_mode=True)
            else:
                outline = outline_for_text(
                    video_title,
                    plain,
                    model=llm_model,
                    api_key=openai_api_key,
                    playlist_mode=True,
                )
            return (video_title, outline, None)
        except Exception as exc:
            return (video_title, None, str(exc))

    results: List[Tuple[str, Optional[OutlineResponse], Optional[str]]] = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(process_entry, e) for e in entries]
        for fut in concurrent.futures.as_completed(futures):
            results.append(fut.result())

    # Preserve original order by mapping back
    title_to_result = {title: (outline, err) for title, outline, err in results}
    for entry in entries:
        video_title = entry.get('title', 'Untitled')
        document.add_heading(video_title, level=2)
        if include_raw_transcript:
            # Intentionally not writing raw transcript as requested; keep disabled by default
            pass
        outline, err = title_to_result.get(video_title, (None, "No result"))
        if err:
            document.add_paragraph(f"LLM summarization failed: {err}")
        elif outline:
            write_outline_to_docx(document, outline)

    filename = os.path.join(output_dir, f"{playlist_title}.docx")
    document.save(filename)
    return filename


def write_outline_to_docx(document: Document, outline: OutlineResponse) -> None:
    """
    Write H2/H3 and bullet content to the DOCX document from outline produced by LLM.
    """
    for section in outline.sections:
        document.add_heading(section.title, level=3)
        for sub in section.subsections:
            document.add_heading(sub.title, level=4)
            for bullet in sub.bullets:
                document.add_paragraph(bullet, style='List Bullet')


def generate_single_video_docx(
    video_url: str,
    output_dir: str = "playlists_docx",
    *,
    use_llm: bool = True,
    llm_model: str = "gpt-4o-mini",
    openai_api_key: Optional[str] = None,
) -> str:
    """
    Create a DOCX for a single video. File name and Heading 1 are the video title.
    Heading 2/3 carry the content outline from the LLM, aligned to the video's flow.
    Transcript is not written into the final document.
    """
    os.makedirs(output_dir, exist_ok=True)
    info = get_video_info(video_url)
    video_title = info.get('title', 'Video')

    document = Document()
    document.add_heading(video_title, level=1)

    transcript_text = get_video_transcript(video_url)
    plain = captions_to_plaintext(transcript_text)

    if use_llm:
        duration = int(info.get('duration') or 0)
        if duration > 1800:
            segments = captions_to_segments(transcript_text, segment_seconds=1800)
            outlines: List[OutlineResponse] = []
            for seg_index, start_sec, seg_text in segments:
                outline_seg = outline_for_text(
                    f"{video_title} — Part {seg_index+1}",
                    seg_text,
                    model=llm_model,
                    api_key=openai_api_key,
                    playlist_mode=False,
                )
                outlines.append(outline_seg)
            outline = merge_outlines(video_title, outlines, model=llm_model, api_key=openai_api_key, playlist_mode=False)
        else:
            outline = outline_for_text(
                video_title,
                plain,
                model=llm_model,
                api_key=openai_api_key,
                playlist_mode=False,
            )
        write_outline_to_docx(document, outline)

    filename = os.path.join(output_dir, f"{video_title}.docx")
    document.save(filename)
    return filename


